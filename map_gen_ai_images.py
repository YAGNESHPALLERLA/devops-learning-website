import docx
import json
from pathlib import Path

# Read the document
doc = docx.Document('public/downloads/Gen Ai.docx')

# Load existing structure
with open('gen_ai_content_structure.json', 'r', encoding='utf-8') as f:
    structure = json.load(f)

# Map images to sections by analyzing document structure
# We'll approximate image positions based on paragraph order
image_mapping = {}
current_section_index = 0
image_counter = 0

for i, para in enumerate(doc.paragraphs):
    # Check if paragraph has images
    for rel in para.part.rels.values():
        if "image" in rel.target_ref:
            image_counter += 1
            # Find which section this image belongs to
            # Look backwards to find the most recent section heading
            section_found = None
            for j in range(i, -1, -1):
                prev_para = doc.paragraphs[j]
                prev_text = prev_para.text.strip()
                if prev_text and (prev_text[0].isdigit() and '.' in prev_text[:5] and len(prev_text) > 10):
                    section_found = prev_text
                    break
            
            if section_found:
                # Find section index
                for idx, section in enumerate(structure['sections']):
                    if section['title'] == section_found:
                        if idx not in image_mapping:
                            image_mapping[idx] = []
                        image_mapping[idx].append(image_counter)
                        break

print(f"Image mapping: {image_mapping}")
print(f"Total images mapped: {sum(len(imgs) for imgs in image_mapping.values())}")

# Save mapping
with open('gen_ai_image_mapping.json', 'w', encoding='utf-8') as f:
    json.dump(image_mapping, f, indent=2)

print("Saved image mapping to gen_ai_image_mapping.json")

